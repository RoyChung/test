#!/usr/bin/env python3
"""Generate HKU TPG New Course Proposal from the Emerging Issues syllabus.
   Uses TPG form as template and fills with syllabus content."""

import os
import re
from docx import Document
from docx.shared import Pt

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROPOSAL_TEMPLATE = os.path.join(PROJECT_ROOT, "Stage", "Tanna", "HKU Law", "TPG - New Course Proposal Form_draft 2026-27 (TC edits)_3.docx")
SYLLABUS_PATH = os.path.join(PROJECT_ROOT, "Stage", "Tanna", "HKU Law", "Emerging_Issues_Global_Compliance_Syllabus.docx")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "Stage", "Tanna", "HKU Law", "TPG_New_Course_Proposal_Emerging_Issues.docx")


# Content to fill the proposal form
PROPOSAL_CONTENT = {
    "teacher": "Tanna Chong",
    "course_name": "Emerging Issues in a Global Compliance Context",
    "prereq": "None",
    "llm_programme": "Compliance and Regulation — Specialized Topics Course",

    "course_description": """This course examines emerging compliance issues in a global context. Regulatory developments in financial crime prevention, digital assets, and cross-border coordination are evolving rapidly across jurisdictions.

Topics include the role of international standard-setting bodies (FSB, BCBS, IOSCO, FATF), global regulatory developments for stablecoins and CBDCs, the Crypto-Asset Reporting Framework (CARF), Travel Rule implementation, AI and RegTech in compliance, and the impact of sanctions and geopolitical fragmentation.

The course combines policy and regulatory analysis with practical considerations for compliance professionals. It comprises 11 three-hour seminars covering international standard-setting, FATF standards, global regulatory development for digital assets, tax and prudential frameworks, law enforcement perspectives, Travel Rule and VASP compliance, digital finance convergence, local implementation, AI/RegTech, and sanctions. The final session is dedicated to group project presentations.

Assessment: Group project presentation (80%); Class participation (20%).""",

    "justification": """Many students in the LLM(CR) programme are industry professionals. This course complements the existing curriculum by offering a global and comparative perspective on emerging compliance issues—including CARF, Travel Rule, DeFi regulation, and AI in compliance—so that students can see how these developments affect their day-to-day work.

Emerging compliance issues are inherently global: they arise and evolve across multiple jurisdictions simultaneously. The course synthesises international standard-setting, jurisdictional implementation, and practical compliance challenges. It complements existing courses—such as LLAW6256 (AML/CFT) and LLAW6254/6255 (Compliance: Regulation in Practice / Financial Markets)—by emphasising: (a) international and comparative regulatory frameworks; (b) emerging topics (CARF, Travel Rule, digital assets, AI/RegTech); and (c) policy and implementation perspectives relevant to practitioners.""",

    "clos": [
        "Have a solid understanding of the functions of FSB, BCBS, IOSCO, and FATF and their relevance to compliance.",
        "Be familiar with global regulatory developments for stablecoins, CBDCs, and comparative frameworks (MiCA, HK, US, Singapore).",
        "Understand CARF and prudential treatment of crypto-asset exposures.",
        "Appreciate Travel Rule requirements and cross-border VASP compliance challenges.",
        "Be able to critically analyse emerging issues such as AI, RegTech, sanctions, and geopolitical fragmentation in a compliance context.",
        "Demonstrate the ability to research and present on an emerging compliance issue in a global context.",
    ],

    # CLO to PLO mapping: A=knowledge, B=apply/analyse, C=practical, D=present, E=ethics, F=social
    "plo_mapping": {
        "CLO 1": ["PLO A"],
        "CLO 2": ["PLO A", "PLO B"],
        "CLO 3": ["PLO A"],
        "CLO 4": ["PLO A", "PLO C"],
        "CLO 5": ["PLO B", "PLO C"],
        "CLO 6": ["PLO B", "PLO D"],  # Research and present
    },

    "assessment": "Group project presentation (80%): Students work in groups to research and present on a selected emerging compliance issue in a global context. Class participation (20%): Attendance and constructive contribution to discussions (pass/fail).",

    "assessment_table": [
        ("Group project presentation", "Final session", "80", "Group review meeting / individual feedback", "CLO 1–6"),
        ("Class participation", "On-course", "20", "General course report", "CLO 1–6"),
    ],

    "teaching_materials": """Core materials from FATF, FSB, Basel Committee, OECD, BIS, and relevant jurisdiction publications (HKMA, MAS, OFAC, EU). Specific readings with URLs provided for each of the 11 sessions. Key sources: FATF Recommendations and guidance on VAs/VASPs; FSB reports on DeFi, stablecoins, cross-border payments; OECD CARF; Basel Committee prudential treatment of crypto-assets; IOSCO thematic reviews.""",

    "learning_activities": """Seminar format (3 hours per session × 11 sessions = 33 hours). Each session combines: (1) lecture on core concepts; (2) discussion of assigned readings; (3) case studies and comparative analysis; (4) Q&A. The final session is devoted to group project presentations. Student participation through discussion is expected and assessed (20% class participation).""",
}


def clear_cell(cell):
    """Clear cell content and set new text."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.clear()
    if not cell.paragraphs:
        cell.add_paragraph()
    return cell.paragraphs[0]


def set_cell_text(cell, text):
    """Set cell text, fully replacing content."""
    cell.text = text


def fill_proposal():
    doc = Document(PROPOSAL_TEMPLATE)
    tables = doc.tables

    # Table 0: Proposed Teacher - row 0 cell 0
    if len(tables) > 0:
        set_cell_text(tables[0].rows[0].cells[0], PROPOSAL_CONTENT["teacher"])

    # Table 1: Proposed Course Name
    if len(tables) > 1:
        set_cell_text(tables[1].rows[0].cells[0], PROPOSAL_CONTENT["course_name"])

    # Table 2: Pre-requisite
    if len(tables) > 2:
        set_cell_text(tables[2].rows[0].cells[0], PROPOSAL_CONTENT["prereq"])

    # Table 3: Course Description (200-400 words)
    if len(tables) > 3:
        set_cell_text(tables[3].rows[0].cells[0], PROPOSAL_CONTENT["course_description"])
    # Table 4: Justification
    if len(tables) > 4:
        set_cell_text(tables[4].rows[0].cells[0], PROPOSAL_CONTENT["justification"])

    # Find and fill paragraph-based sections (Course Description, Justification, etc.)
    # These appear after specific headings. We need to insert content after the heading paragraph.
    # Strategy: Find the paragraph that contains "Course Description (200-400 words)" and the next empty one
    # or the one that says "Please provide a brief course description"
    found_desc = False
    found_just = False
    found_clos = False
    found_assess = False
    found_materials = False
    found_activities = False

    # Justification: might be in a table. Try Table 4 = Course Desc, check if another table for Justification.
    # Form structure: Tables 0-3 = teacher, course, prereq, LLM. Then Course Description (200-400 words) 
    # could be in a large cell - Table 4. Justification might be in Table 5 - but Table 5 has CLO text.
    # Alternative: Table 3 = Course Desc, Table 4 = Justification. Try both in tables 4 and 5.
    # Table 5 row 0 has "At the end of this course..." - so Table 5 is CLO. No justification table.
    # Justification filled in Table 4 above.

    # Fill Assessment - find instruction and fill next paragraph
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if "Briefly describe the types of assessment activities" in t and i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            next_t = (next_p.text or "").strip()
            if len(next_t) < 150:  # Short or empty
                next_p.clear()
                next_p.add_run(PROPOSAL_CONTENT["assessment"])
                found_assess = True
            break

    # Fill CLO table (Table 6) - add row if needed for 6 CLOs
    if len(tables) > 6:
        clo_table = tables[6]
        num_clos = len(PROPOSAL_CONTENT["clos"])
        while len(clo_table.rows) < num_clos:
            clo_table.add_row()
        for i, clo in enumerate(PROPOSAL_CONTENT["clos"]):
            if i < len(clo_table.rows):
                set_cell_text(clo_table.rows[i].cells[0], f"CLO {i + 1}")
                if len(clo_table.rows[i].cells) > 1:
                    set_cell_text(clo_table.rows[i].cells[1], clo)

    # Fill PLO mapping table (Table 7) - rows 1+ for CLO 1-6, cols for PLO A-F
    if len(tables) > 7:
        plo_table = tables[7]
        plo_cols = {"PLO A": 1, "PLO B": 2, "PLO C": 3, "PLO D": 4, "PLO E": 5, "PLO F": 6}
        num_clos = len(PROPOSAL_CONTENT["clos"])
        while len(plo_table.rows) <= num_clos:
            plo_table.add_row()
        for i in range(num_clos):
            row_idx = i + 1
            if row_idx < len(plo_table.rows):
                set_cell_text(plo_table.rows[row_idx].cells[0], f"CLO {i + 1}")
                plos = PROPOSAL_CONTENT["plo_mapping"].get(f"CLO {i + 1}", [])
                for plo in plos:
                    if plo in plo_cols:
                        col_idx = plo_cols[plo]
                        if col_idx < len(plo_table.rows[row_idx].cells):
                            set_cell_text(plo_table.rows[row_idx].cells[col_idx], "X")

    # Fill assessment table (Table 9)
    if len(tables) > 9:
        assess_table = tables[9]
        for i, row_data in enumerate(PROPOSAL_CONTENT["assessment_table"]):
            if i + 1 < len(assess_table.rows):
                row = assess_table.rows[i + 1]
                for j, val in enumerate(row_data):
                    if j < len(row.cells):
                        set_cell_text(row.cells[j], str(val))

    # Fill coursework/exam ratio table (Table 8)
    if len(tables) > 8:
        ratio_table = tables[8]
        set_cell_text(ratio_table.rows[0].cells[1], "100")
        set_cell_text(ratio_table.rows[0].cells[2], "Coursework")
        set_cell_text(ratio_table.rows[0].cells[4], "0")

    # Teaching materials and learning activities - often in tables 11, 12
    if len(tables) > 11:
        set_cell_text(tables[11].rows[0].cells[0], PROPOSAL_CONTENT["teaching_materials"])
    if len(tables) > 12:
        # Table 12: Learning activities - typically first data cell
        set_cell_text(tables[12].rows[0].cells[1], PROPOSAL_CONTENT["learning_activities"])

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = fill_proposal()
    print(f"Exported to: {out}")
