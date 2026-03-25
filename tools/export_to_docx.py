#!/usr/bin/env python3
"""Export VerifyVASP LinkedIn post and news article to Word document with images."""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def add_image(doc, path, width=5.5, caption=None):
    """Add image to document with optional caption."""
    if not os.path.exists(path):
        doc.add_paragraph(f"[Image placeholder: {os.path.basename(path)}]")
        return
    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(path, width=Inches(width))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = "Caption"
    except Exception as e:
        doc.add_paragraph(f"[Could not embed image: {e}]")


def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading("VerifyVASP Content: FATF Report on Offshore VASPs", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # --- LinkedIn Post ---
    doc.add_heading("1. LinkedIn Post (Draft)", level=1)
    
    linkedin_content = """
📢 **FATF's new report shines a spotlight on Offshore VASPs – and why it matters for the industry**

The Financial Action Task Force (FATF) has just released "Understanding and Mitigating the Risks of Offshore Virtual Asset Service Providers (oVASPs)" – and it highlights serious gaps in oversight that criminals are actively exploiting.

**What are oVASPs?**  
Offshore VASPs are Virtual Asset Service Providers created under the laws of one jurisdiction that provide services to clients in another, often without being licensed or registered where they operate. This can create blind spots for fraud, money laundering, and terrorism financing.

**Key findings:**
• **46%** of jurisdictions still lack an activity-based approach to regulate oVASPs  
• Criminals use oVASPs to launder proceeds from scam compounds and support terrorist groups  
• **Nested relationships** – where unlicensed oVASPs pose as retail customers to access regulated VASPs – are a major vulnerability  
• Fragmented Travel Rule implementation causes blind spots that undermine investigations  

**The way forward** – The report urges jurisdictions and the private sector to:
✓ Adopt activity-based licensing to bring offshore providers under supervision  
✓ Strengthen FIU-to-FIU and supervisor-to-supervisor cooperation  
✓ Enforce sanctions for non-compliance  

At VerifyVASP, we are committed to supporting this global effort through **objective-based Travel Rule compliance**, **VerifyWallet** for self-hosted wallet verification, and **Law Enforcement Solutions** that combine blockchain data with verified transactional intelligence.

Strong compliance and international cooperation are essential as virtual assets move across borders in seconds. We welcome the FATF's guidance and remain dedicated to advancing transparency and security across the virtual asset ecosystem.

🔗 Read the full report: https://www.fatf-gafi.org/content/fatf-gafi/en/publications/Virtualassets/Understanding-Mitigating-Risks-Offshore-VASPs.html

#FATF #AML #CFT #VASPs #TravelRule #VerifyVASP #VirtualAssets #Compliance
""".strip()
    
    def add_formatted_paragraph(doc, text):
        p = doc.add_paragraph()
        parts = text.replace("**", "\x00").split("\x00")
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1:
                run.bold = True
        p.paragraph_format.space_after = Pt(6)

    for para in linkedin_content.split("\n\n"):
        if para.strip():
            add_formatted_paragraph(doc, para)
    
    doc.add_paragraph()
    doc.add_paragraph("Suggested images for LinkedIn:", style="Heading 3")
    
    img_dir = PROJECT_ROOT
    images = [
        ("1773222174833.png", "Image 1: FATF infographic – What are Offshore VASPs?"),
        ("1773222174833 2.png", "Image 2: FATF report / risk typologies"),
        ("1773222174833 3.png", "Image 3: Additional FATF visuals"),
    ]
    for fname, cap in images:
        path = os.path.join(img_dir, fname)
        add_image(doc, path, width=5.0, caption=cap)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # --- News Article ---
    doc.add_heading("2. News Article (Draft) – VerifyVASP Website Style", level=1)
    doc.add_paragraph("FATF Report on Offshore VASPs: Understanding and Mitigating Risks", style="Title")
    doc.add_paragraph("March 11, 2026 | 4 min read")
    doc.add_paragraph()
    
    news_sections = [
        ("Introduction", """
The Financial Action Task Force (FATF) has released its March 2026 report, "Understanding and Mitigating the Risks of Offshore Virtual Asset Service Providers (oVASPs)", focusing on how gaps in oversight of offshore VASPs are exploited to facilitate large-scale fraud, money laundering, and terrorism financing. The report also presents good practices to detect, license or register, and supervise oVASPs, as well as sanction non-compliant ones.
"""),
        ("Key Findings: The Offshore VASP Landscape", """
**What are Offshore VASPs?**

Offshore VASPs are VASPs created under the laws of one jurisdiction ("home jurisdiction") with or without a physical presence, that provide services to clients domiciled or residing in jurisdictions outside their jurisdiction of incorporation or physical location ("host jurisdiction"). Many operate without being licensed or registered where they actively provide services – creating significant AML/CFT blind spots.

**The Core Problem**

Criminals exploit gaps in regulation and supervision to use offshore VASPs for money laundering, terrorist financing, and proliferation financing. The FATF identifies five key risk typologies:

1. **Active targeting** – Unlicensed or unregistered offshore VASPs targeting users in jurisdictions where they are not regulated
2. **Pooling of customers** – Through offshore group entities rather than locally supervised VASPs, blurring accountability
3. **Divergent approaches** – Across jurisdictions, creating regulatory and supervisory gaps
4. **Fragmented Travel Rule implementation** – Causing blind spots that undermine preventive measures and slow investigations
5. **Nested VASP activity** – When one VASP conducts services through another, obscuring underlying offshore activity where controls are inadequate
"""),
        ("The Illicit Finance Reality", """
The report highlights how oVASPs have been used to:

• **Convert illicit proceeds from scam compounds** – Nigeria's FIU identified how oVASPs and opaque corporate structures facilitated large-scale fraud, with one global VASP-linked wallet holding approximately **USD 600 million** at the time of analysis
• **Provide financial support to terrorist groups** – Indonesia's FIU identified VA-based financial support to terrorist groups in Syria, with oVASPs used for anonymity, conversion between asset types, and rapid layering before funds moved to unhosted wallets
• **Enable regulatory arbitrage** – After India's 2022 VA tax regime, a significant proportion of trading traffic moved from onshore to offshore unregistered VASPs that marketed minimal KYC and encouraged VPN use
"""),
        ("Good Practices for Jurisdictions and the Private Sector", """
**For jurisdictions:**
• Detecting and licensing/registering oVASPs using an activity-based approach
• Enforcing sanctions for non-compliance with AML/CFT/CPF obligations
• Building shared understanding through inter-agency task forces and public-private partnerships
• Maximising supervisor-to-supervisor and FIU-to-FIU cooperation

**For the private sector:**
• Assessing exposure to unlicensed or unregistered oVASPs
• Applying clear, consistent AML/CFT/CPF rules across all group entities
• Ensuring no group entity operates as an oVASP abroad outside regulatory oversight
• Refraining from establishing or maintaining business relationships with unlicensed or unregistered providers
"""),
        ("Jointly Mitigating These Risks & VerifyVASP's Role", """
The FATF emphasises that all VASPs should adopt risk-mitigating measures in accordance with its recommendations. VerifyVASP has and remains committed to a public-private framework:

1. **VerifyWallet** – A scalable solution that VASPs can implement to verify that their customers have control over self-hosted wallets using key signing technologies within the VASP's own secured environment.
2. **Law Enforcement Solutions** – An AI-driven tracing tool offered to the public sector (Law Enforcement Agencies, Financial Intelligence Units, and Regulators) for effective and quick tracing by combining blockchain data with VerifyVASP's verified transactional information accumulated through objective-based Travel Rule compliance.
3. **Public Private Partnerships (PPP)** – VerifyVASP is an active participant in various PPPs established internationally and regionally to jointly combat illicit activity in the industry.

At VerifyVASP, we are committed to advancing this collaborative approach through technical breakthroughs and industry cooperation. Our goal is to strengthen transparency, compliance, and security in virtual asset transactions.
"""),
    ]
    
    for heading, text in news_sections:
        doc.add_heading(heading, level=2)
        for para in text.strip().split("\n\n"):
            p = doc.add_paragraph()
            # Simple bold handling
            parts = para.replace("**", "\x00").split("\x00")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
            p.paragraph_format.space_after = Pt(6)
        doc.add_paragraph()
    
    doc.add_paragraph("Read the full FATF report: https://www.fatf-gafi.org/content/fatf-gafi/en/publications/Virtualassets/Understanding-Mitigating-Risks-Offshore-VASPs.html")
    doc.add_paragraph()
    
    doc.add_heading("Images for News Article", level=2)
    for fname, cap in images:
        path = os.path.join(img_dir, fname)
        add_image(doc, path, width=5.0, caption=cap)
        doc.add_paragraph()
    
    output_path = os.path.join(PROJECT_ROOT, "VerifyVASP_oVASPs_Content.docx")
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    out = create_document()
    print(f"Exported to: {out}")
