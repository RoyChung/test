#!/usr/bin/env python3
"""Generate HKU LLM Course Syllabus for Emerging Issues in a Global Compliance Context.
   Format follows HKU Sample Course Syllabus: Times New Roman, 16pt title, 11pt body, bold section headings."""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SAMPLE_PATH = os.path.join(PROJECT_ROOT, "Stage", "Tanna", "HKU Law", "Sample Course Syllabus.docx")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "Stage", "Tanna", "HKU Law", "Emerging_Issues_Global_Compliance_Syllabus.docx")


LESSONS = [
    (
        "Introduction: Emerging Issues in a Global Compliance Context; International Standard-Setting Bodies (FSB, BCBS, IOSCO)",
        "This session introduces the course framework and the intersection of emerging compliance issues with global regulatory developments. We will explore why emerging issues are inherently global, how regulatory fragmentation and convergence play out across jurisdictions, and what this means for compliance professionals. The session also examines the Financial Stability Board (FSB), Basel Committee on Banking Supervision (BCBS), and International Organization of Securities Commissions (IOSCO)—their mandates, standard-setting processes, and how they coordinate to address systemic risks and cross-border financial regulation.",
        [
            "FATF, High-level Recommendations for the Regulation, Supervision and Oversight of Crypto-asset Activities and Markets (2023), available at: https://www.fatf-gafi.org/content/dam/fatf-gafi/recommendations/2024-Targeted-Update-VA-VASP.pdf.coredownload.inline.pdf",
            "FSB, The Financial Stability Risks of Decentralised Finance (2023), available at: https://www.fsb.org/2023/02/the-financial-stability-risks-of-decentralised-finance/",
            "FSB, Promoting Global Financial Stability: 2024 FSB Annual Report (PDF), available at: https://www.fsb.org/uploads/P181124-2.pdf",
            "Basel Committee, Basel III: international regulatory framework for banks (overview and standards), available at: https://www.bis.org/bcbs/basel3.htm",
            "IOSCO, Thematic Review on Technological Challenges to Effective Market Surveillance (2025), available at: https://www.iosco.org/library/pubdocs/pdf/IOSCOPD809.pdf",
        ],
    ),
    (
        "FATF: Standards, Functions and Jurisdictional Implementation",
        "This session covers the Financial Action Task Force (FATF), its recommendations, and how jurisdictions implement them. Topics include mutual evaluations, the FATF Global Network, beneficial ownership reforms, and payment transparency standards.",
        [
            "FATF, The FATF Recommendations (consolidated PDF), available at: https://www.fatf-gafi.org/content/dam/fatf-gafi/recommendations/FATF%20Recommendations%202012.pdf.coredownload.pdf",
            "FATF, Methodology for assessing compliance with the FATF Recommendations and the effectiveness of AML/CFT systems (2023), available at: https://www.fatf-gafi.org/content/dam/fatf-gafi/publications/Methodology/2023-02-FATF-Methodology.pdf.coredownload.pdf",
            "FATF, Mutual Evaluation of Hong Kong (2024), available at: https://www.fatf-gafi.org/content/fatf-gafi/en/publications/Mutual-evaluations/Hong-kong-2024.html",
        ],
    ),
    (
        "Global Regulatory Development: Stablecoins, CBDCs and Comparative Frameworks",
        "We will compare regulatory approaches to stablecoins and central bank digital currencies (CBDCs) across MiCA (EU), Hong Kong, the US, and Singapore. The session explores convergence and divergence in licensing, reserve requirements, and consumer protection.",
        [
            "FSB, High-level Recommendations for the Regulation, Supervision and Oversight of Global Stablecoin Arrangements (2023), available at: https://www.fsb.org/uploads/P170723-3.pdf",
            "BIS, Central bank digital currencies: foundational principles and core features (2020), available at: https://www.bis.org/publ/othp33.pdf",
            "HKMA, Consultation Conclusions on legislative proposal for stablecoin regulation (2024), available at: https://www.hkma.gov.hk/media/eng/doc/key-information/consultations/20240719e1.pdf",
        ],
    ),
    (
        "CARF and Prudential Treatment: Tax Transparency and Crypto-Asset Exposures",
        "This session examines the OECD Crypto-Asset Reporting Framework (CARF) and the Basel Committee's prudential treatment of crypto-asset exposures. We will discuss automatic exchange of information, due diligence for crypto-asset intermediaries, and capital requirements for banks' crypto exposures.",
        [
            "OECD, Crypto-Asset Reporting Framework and Amendments to the Common Reporting Standard (2022), available at: https://www.oecd.org/tax/exchange-of-tax-information/crypto-asset-reporting-framework-and-amendments-to-the-common-reporting-standard.pdf",
            "Basel Committee, Prudential treatment of cryptoasset exposures (2022), available at: https://www.bis.org/bcbs/publ/d545.pdf",
            "OECD, Crypto-Asset Reporting Framework: FAQs and implementation update (2025), available at: https://www.oecd.org/content/dam/oecd/en/topics/policy-issues/tax-transparency-and-international-co-operation/faqs-crypto-asset-reporting-framework.pdf",
        ],
    ),
    (
        "Law Enforcement, Financial Crime and Compliance Preparedness",
        "We will explore the law enforcement and financial crime perspective on emerging compliance challenges. Topics include typologies of illicit finance, investigative approaches, and how compliance functions can better support detection, reporting, and cooperation with authorities.",
        [
            "FATF, Money Laundering and Terrorist Financing Red Flags (2023), available at: https://www.fatf-gafi.org/content/dam/fatf-gafi/guidance/RBA-Financial-Inclusion-Red-Flags.pdf.coredownload.pdf",
            "FATF, Illicit Financial Flows from Cyber-Enabled Fraud and Cyber-Enabled Theft (2024), available at: https://www.fatf-gafi.org/en/publications/Methods-and-trends/Cyber-enabled-fraud-and-cyber-enabled-theft.html",
            "Europol, European Financial and Economic Crime Threat Assessment 2024, available at: https://www.europol.europa.eu/publication-events/main-reports/european-financial-and-economic-crime-threat-assessment-2024",
        ],
    ),
    (
        "Travel Rule and Cross-Border VASP Compliance",
        "This session focuses on FATF Recommendation 16 (Travel Rule) and its application to virtual asset service providers (VASPs). We will discuss cross-border information sharing, implementation challenges, and infrastructure solutions for Travel Rule compliance.",
        [
            "FATF, Updated Guidance for a Risk-Based Approach to Virtual Assets and VASPs (2021), available at: https://www.fatf-gafi.org/content/dam/fatf-gafi/guidance/Updated-Guidance-VA-VASP.pdf.coredownload.inline.pdf",
            "FATF, Understanding and Mitigating the Risks of Offshore VASPs (2026), available at: https://www.fatf-gafi.org/content/fatf-gafi/en/publications/Virtualassets/Understanding-Mitigating-Risks-Offshore-VASPs.html",
            "FATF, Best Practices on Travel Rule Supervision (2025), available at: https://www.fatf-gafi.org/en/publications/Virtualassets/Travel-Rule-Best-Practices.html",
        ],
    ),
    (
        "Digital Finance Meets Traditional Finance: Payment Regulation and Convergence",
        "We will examine how digital finance interfaces with traditional finance, including payment regulation, licensing regimes, and the blurring lines between banks, payment institutions, and digital asset providers. Topics include open banking, instant payments, and regulatory sandboxes.",
        [
            "BIS CPMI, Central bank digital currencies (CBDCs) – system design and interopability (2024), available at: https://www.bis.org/cpmi/publ/d221.pdf",
            "FSB, G20 Roadmap for Enhancing Cross-border Payments – Progress report (2024), available at: https://www.fsb.org/uploads/P221024-3.pdf",
            "HKMA, Fintech Supervisory Sandbox – Guide for Applicants, available at: https://www.hkma.gov.hk/media/eng/doc/key-functions/fintech/20240422e1.pdf",
        ],
    ),
    (
        "Local Legislative Implementation: From Global Standards to Domestic Regulation",
        "This session discusses how global standards are transposed into domestic law and regulation. We will compare implementation approaches, timing, and the role of regulators in interpreting and enforcing international standards.",
        [
            "FATF, Methodology for assessing compliance with the FATF Recommendations (2023) (PDF), available at: https://www.fatf-gafi.org/content/dam/fatf-gafi/publications/Methodology/2023-02-FATF-Methodology.pdf.coredownload.pdf",
            "FSI (BIS), Implementation of Basel III – Executive summaries, available at: https://www.bis.org/fsi/fsisummaries/basel_framework.htm",
            "FATF, Jurisdictional Mutual Evaluation Reports (select jurisdictions), available at: https://www.fatf-gafi.org/en/countries.html",
        ],
    ),
    (
        "AI, RegTech and SupTech in Compliance",
        "We will explore the use of artificial intelligence, regulatory technology (RegTech), and supervisory technology (SupTech) in compliance. Topics include transaction monitoring, KYC automation, algorithmic bias, and regulatory expectations for AI governance.",
        [
            "Bank of England / FCA, Machine Learning in UK Financial Services (2019), available at: https://www.bankofengland.co.uk/-/media/boe/files/fintech/machine-learning-in-uk-financial-services.pdf",
            "MAS, Veritas 2.0 – Feasibility Study on the use of AI and Data Analytics (2021), available at: https://www.mas.gov.sg/-/media/mas-media-lib/development/fintech/veritas/veritas-20-feasibility-study.pdf",
            "FATF, Opportunities and Challenges of New Technologies for AML/CFT (2021), available at: https://www.fatf-gafi.org/content/dam/fatf-gafi/publications/MethodsAndTrends/Opportunities-Challenges-New-Technologies-AML-CFT.pdf.coredownload.pdf",
        ],
    ),
    (
        "Sanctions, Geopolitics and Regulatory Fragmentation",
        "This session examines sanctions regimes, geopolitical tensions, and their impact on compliance. We will discuss multi-jurisdictional sanctions (OFAC, EU, UK, HK), fragmentation risks, and how compliance teams navigate conflicting obligations.",
        [
            "OFAC, A Framework for OFAC Compliance Commitments (2019), available at: https://ofac.treasury.gov/media/176991/download",
            "EU Commission, Restrictive measures (sanctions) – consolidated list, available at: https://finance.ec.europa.eu/eu-tax-and-customs-union/eu-sanctions_en",
            "Hong Kong, United Nations (Anti-Terrorism Measures) Ordinance (Cap. 575), available at: https://www.elegislation.gov.hk/hk/cap575",
        ],
    ),
    (
        "Synthesis: Group Project Presentations",
        "The final session is dedicated to group project presentations. Students will present their research on a selected emerging compliance issue in a global context, synthesising themes from the course.",
        None,
    ),
]

# Font constants per Sample Course Syllabus
FONT_NAME = "Times New Roman"
FONT_SIZE_TITLE = Pt(16)
FONT_SIZE_BODY = Pt(11)


def set_run_font(run, size=Pt(11), bold=False, name="Times New Roman"):
    """Set font properties on a run to match Sample format."""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold


def add_para(doc, text, bold=False, style="Normal", alignment=None, space_before=None):
    """Add paragraph with Times New Roman 11pt. Returns the paragraph."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if alignment:
        p.alignment = alignment
    r = p.add_run(text)
    set_run_font(r, size=FONT_SIZE_BODY, bold=bold)
    return p


def add_title(doc, text):
    """Add centered title: Times New Roman 16pt, bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=FONT_SIZE_TITLE, bold=True)
    return p


def add_section_heading(doc, text, space_before=Pt(8)):
    """Add section heading: Times New Roman 11pt, bold. Slight space before for visual separation."""
    return add_para(doc, text, bold=True, space_before=space_before)


def add_body(doc, text):
    """Add body paragraph: Times New Roman 11pt."""
    return add_para(doc, text)


def add_list_item(doc, text):
    """Add list item using List Paragraph style: Times New Roman 11pt."""
    return add_para(doc, text, style="List Paragraph")


def create_syllabus():
    # Use Sample as template to inherit styles, fonts, colors, section settings
    doc = Document(SAMPLE_PATH)

    # Clear existing body content (keeps styles, section, document properties)
    body = doc.element.body
    for child in list(body):
        body.remove(child)

    # Title
    add_title(doc, "Emerging Issues in a Global Compliance Context")
    add_para(doc, "(LLM) — Faculty of Law, The University of Hong Kong")
    doc.add_paragraph()

    # Course Description
    add_section_heading(doc, "Course Description")
    add_body(
        doc,
        "This course examines emerging compliance issues in a global context. Regulatory developments in financial crime prevention, digital assets, and cross-border coordination are evolving rapidly across jurisdictions. Topics include the role of international standard-setting bodies (FSB, BCBS, IOSCO, FATF), global regulatory developments for stablecoins and CBDCs, the Crypto-Asset Reporting Framework (CARF), Travel Rule implementation, AI and RegTech in compliance, and the impact of sanctions and geopolitical fragmentation. The course combines policy and regulatory analysis with practical considerations for compliance professionals.",
    )
    doc.add_paragraph()

    # Key Dates
    add_section_heading(doc, "Key Dates:")
    add_body(doc, "[To be inserted]")
    doc.add_paragraph()

    # Course Overview
    add_section_heading(doc, "Course Overview")
    add_section_heading(doc, "Aims and Objectives")
    add_body(doc, "The key objectives of the course are as follows:")
    doc.add_paragraph()
    objectives = [
        "1. Comprehensive Understanding: To provide students with a comprehensive understanding of emerging compliance issues in a global context.",
        "2. Regulatory Knowledge: To equip students with knowledge of international standard-setting bodies and how their standards are implemented across jurisdictions.",
        "3. Policy and Practice: To examine key regulatory developments in digital assets, tax transparency (CARF), Travel Rule, and related areas.",
        "4. Technology and Compliance: To explore the intersection of technology (AI, RegTech) with compliance and regulatory supervision.",
        "5. Critical Analysis: To encourage critical analysis of regulatory convergence, divergence, and fragmentation in an era of geopolitical tension.",
    ]
    for obj in objectives:
        add_body(doc, obj)
    doc.add_paragraph()

    # Teaching Methodology
    add_section_heading(doc, "Teaching / Learning Methodology / Methods")
    add_body(
        doc,
        "The course is organised in a seminar format (3 hours per session). Student participation through discussion and questions is encouraged and forms part of the learning experience. Each session combines lecture, case studies, and interactive discussion.",
    )
    doc.add_paragraph()

    # Course Organisation
    add_section_heading(doc, "Course Organisation / Coverage")
    add_body(
        doc,
        "The course comprises 11 sessions covering international standard-setting (FSB, BCBS, IOSCO), global regulatory development for digital assets, tax and prudential frameworks, Travel Rule and VASP compliance, digital finance convergence, local implementation, AI/RegTech, and sanctions. The final session is dedicated to group project presentations.",
    )
    doc.add_paragraph()

    # Course Outcomes
    add_section_heading(doc, "Course Outcomes")
    add_body(doc, "At the end of the course, students should:")
    outcomes = [
        "Have a solid understanding of the functions of FSB, BCBS, IOSCO, and FATF and their relevance to compliance.",
        "Be familiar with global regulatory developments for stablecoins, CBDCs, and comparative frameworks (MiCA, HK, US, Singapore).",
        "Understand CARF and prudential treatment of crypto-asset exposures.",
        "Appreciate Travel Rule requirements and cross-border VASP compliance challenges.",
        "Be able to critically analyse emerging issues such as AI, RegTech, sanctions, and geopolitical fragmentation in a compliance context.",
        "Demonstrate the ability to research and present on an emerging compliance issue in a global context.",
    ]
    for out in outcomes:
        add_list_item(doc, out)
    doc.add_paragraph()

    # Assessment
    add_section_heading(doc, "Assessment")
    add_body(doc, "The course assessment is based on:")
    add_body(
        doc,
        "Group project presentation (80%): Students will work in groups to research and present on a selected emerging compliance issue in a global context. The presentation will be delivered in the final session. Further details on group formation, topic selection, and presentation format will be provided at the start of the course.",
    )
    add_body(
        doc,
        "Class participation (20%): This comprises class attendance and constructive contribution to discussions. Graded on a pass/fail basis.",
    )
    add_body(
        doc,
        "The Law Faculty's rules prohibiting cheating, plagiarism and taking unfair advantage apply strictly. All written work will be run through plagiarism detection software as part of the submission process where applicable.",
    )
    doc.add_paragraph()

    # Reading List
    add_section_heading(doc, "Reading list")
    add_body(
        doc,
        "Core materials will be drawn from FATF, FSB, Basel Committee, OECD, and relevant jurisdiction publications. Links and specific readings will be provided for each session.",
    )
    for item in [
        "FATF: https://www.fatf-gafi.org",
        "FSB: https://www.fsb.org",
        "Basel Committee: https://www.bis.org/bcbs",
        "OECD (CARF): https://www.oecd.org/tax/automatic-exchange/",
    ]:
        add_list_item(doc, item)
    doc.add_paragraph()

    # Attendance
    add_section_heading(doc, "Attendance")
    add_body(
        doc,
        "Regular class attendance and participation is expected. If you cannot attend a scheduled session due to illness or other matters, please contact [*].",
    )
    doc.add_paragraph()

    # Preparation
    add_section_heading(doc, "Preparation")
    add_body(doc, "Students should read the assigned materials before each session.")
    doc.add_paragraph()

    # Course Syllabus
    add_section_heading(doc, "Course Syllabus")
    doc.add_page_break()

    for i, (title, desc, reading) in enumerate(LESSONS, 1):
        add_section_heading(doc, f"Seminar No. {i}:  {title}")
        doc.add_paragraph()
        add_section_heading(doc, "[Date: To be inserted]")
        doc.add_paragraph()
        add_body(doc, desc)
        doc.add_paragraph()
        add_section_heading(doc, "Recommended readings:")
        doc.add_paragraph()
        if reading:
            for item in reading:
                add_list_item(doc, item)
        else:
            add_body(doc, "N/A — Presentation session.")
        doc.add_paragraph()
        doc.add_paragraph()

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = create_syllabus()
    print(f"Exported to: {out}")
